import os
import docx
import re

dictionary_of_features = {}
def docx2txt(dirname='/home/andrey/datasets/doc_train/train/14'):
    array_of_path_docx = []
    for i in os.listdir(dirname):
        if re.fullmatch( '[\s.\w]+.docx', i) is not None:
            array_of_path_docx.append(i)
    print(array_of_path_docx)

    text = []
    for i in array_of_path_docx:
        doc = docx.Document(os.path.join(dirname, i))
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
#print('\n\n'.join(text))
    return doc, text
doc, text = docx2txt()

#Добыча нефти
OBJECT1 = re.compile('УЭЦН')
OBJECT2 = re.compile('однотруб[\w+] герметизир[\w+]')
for paragraph in text:
    if re.search(OBJECT1, paragraph):
        dictionary_of_features['Способ добычи'] = 'ЭЦН'
    elif re.search(OBJECT2, paragraph):
        dictionary_of_features['Способ добычи'] = 'напорная однотрубная герметизированная система сбора'
    else:
        dictionary_of_features['Способ добычи'] = 'ЭЦН'


#Абсолютный мининум
OBJECT3 = re.compile('[аА]бсолют[\w\s]+миним[–\s\w,]+[.,\d]+\s*[°]*')
for paragraph in text:
    if re.search(OBJECT3, paragraph):
        sentense = re.findall(OBJECT3, paragraph)
        dictionary_of_features['Абсолютный минимум'] = \
            float('-'+''.join(re.findall(r'[\d]+[.,]*[\d]*', ' '.join(sentense))).replace(',', '.'))
        break
    else:
        dictionary_of_features['Абсолютный минимум'] = None
#Абсолютный максимум
OBJECT4 = re.compile('[аА]бсолют[\w\s]+макс[-\s\w]+[+.,\d]+\s*[°]*')
for paragraph in text:
    if re.search(OBJECT4, paragraph):
        sentense = re.findall(OBJECT4, paragraph)
        print(sentense)
        dictionary_of_features['Абсолютный максимум'] = \
            float(''.join(re.findall(r'[\d]+[.,]*[\d]*', ' '.join(sentense))).replace(',', '.'))
        break
    else:
        dictionary_of_features['Абсолютный максимум'] = None

#Среднемесячная температура самого холодного месяца
OBJECT5 = re.compile('холодн[–\s\w,]+месяц[–\s\w,]+[.,\d]+\s*[°]*')
for paragraph in text:
    if re.search(OBJECT5, paragraph):
        sentense = re.findall(OBJECT5, paragraph)
        dictionary_of_features['Среднемесячная температура самого холодного месяца'] = \
            float('-'+''.join(re.findall(r'[\d]+[.,]*[\d]*', ' '.join(sentense))).replace(',', '.'))
        break
    else:
        dictionary_of_features['Среднемесячная температура самого холодного месяца'] = None

#Средняя температура наиболее холодной пятидневки
OBJECT6 = re.compile('холодн[–\s\w,]+пятидневк[–\w\s]+[.,\d]+\s*[°]*')
for paragraph in text:
    if re.search(OBJECT6, paragraph):
        sentense = re.findall(OBJECT6, paragraph)
        dictionary_of_features['Средняя температура наиболее холодной пятидневки'] = \
            float('-'+''.join(re.findall(r'[\d]+[.,]*[\d]*', ' '.join(sentense))).replace(',', '.'))
        break
    else:
        dictionary_of_features['Средняя температура наиболее холодной пятидневки'] = None

#
print('\n\n'.join(text))
print(dictionary_of_features)















tables = []
for table in doc.tables:
    tables.append(table)

'''''
for table in tables:
    for row in table.rows:
        print([cell.text for cell in row.cells])
'''
